"""Shared .docx section builders. Every profile renders through these same
functions -- the only thing that varies between resumes is the filtered data
passed in, never the styling code."""
from pathlib import Path

from docx import Document

TEMPLATE_PATH = Path(__file__).resolve().parent.parent.parent / "templates" / "resume_style.docx"


def new_document() -> Document:
    return Document(str(TEMPLATE_PATH))


def add_header(doc, contact):
    doc.add_paragraph(contact.get("name", ""), style="NameHeader")
    parts = [p for p in (
        contact.get("phone"), contact.get("email"),
        contact.get("location"), contact.get("linkedin"),
    ) if p]
    doc.add_paragraph(" | ".join(parts), style="ContactLine")


def add_headline(doc, headline):
    if headline:
        doc.add_paragraph(headline, style="Headline")


def add_summary(doc, summary):
    if summary and summary.strip():
        doc.add_paragraph(" ".join(summary.split()), style="PlainText")


def _add_section_heading(doc, text):
    doc.add_paragraph(text.upper(), style="SectionHeading")


def _format_date_range(start, end):
    end_label = "Present" if end == "present" else end
    return f"{start} – {end_label}"


def _add_job(doc, job, bullets):
    p = doc.add_paragraph(style="JobTitle")
    p.add_run(f"{job.title} — {job.employer}")

    meta_bits = [b for b in (job.location, _format_date_range(job.start_date, job.end_date)) if b]
    doc.add_paragraph(" | ".join(meta_bits), style="JobMeta")

    for bullet in bullets:
        doc.add_paragraph(bullet.text, style="BulletList")


def add_experience_section(doc, jobs_with_bullets):
    if not jobs_with_bullets:
        return
    _add_section_heading(doc, "Experience")
    for job, bullets in jobs_with_bullets:
        _add_job(doc, job, bullets)


def add_skills_section(doc, categories):
    if not categories:
        return
    _add_section_heading(doc, "Skills")
    for category in categories:
        skills = category.get("skills", [])
        if not skills:
            continue
        p = doc.add_paragraph(style="SkillCategory")
        p.add_run(f"{category['name']}: ")
        run = p.add_run(", ".join(skills))
        run.bold = False


def add_education_section(doc, education):
    if not education:
        return
    _add_section_heading(doc, "Education")
    for entry in education:
        if entry.get("degree"):
            line = f"{entry['degree']} in {entry['field']}, {entry['institution']}"
        else:
            line = f"{entry['field']}, {entry['institution']}"
        if entry.get("location"):
            line += f" — {entry['location']}"
        if entry.get("honors"):
            line += f", {entry['honors']}"
        doc.add_paragraph(line, style="PlainText")


def add_certifications_section(doc, certifications):
    if not certifications:
        return
    _add_section_heading(doc, "Certifications")
    for cert in certifications:
        line = cert["name"]
        if cert.get("issuer"):
            line += f" — {cert['issuer']}"
        doc.add_paragraph(line, style="PlainText")


def add_awards_section(doc, awards):
    if not awards:
        return
    _add_section_heading(doc, "Awards")
    for award in awards:
        line = award["name"]
        if award.get("context"):
            line += f" — {award['context']}"
        doc.add_paragraph(line, style="PlainText")
